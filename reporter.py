from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT,WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np
from docx.oxml.shared import qn
from docx.oxml.xmlchemy import OxmlElement
import matplotlib.pyplot as plt
from draw import draw
import docx2pdf







def set_table_header_bg_color(tc):
    """
    set background shading for Header Rows
    """
    tblCellProperties = tc._element.tcPr
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), "d3d3d3")  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tblCellProperties.append(clShading)


def merge_cells_by_column(table, column_index):
    cells = [row.cells[column_index] for row in table.rows]

    for i in range(len(cells) - 1):
        if cells[i].text == cells[i + 1].text:
            cells[i + 1].merge(cells[i])


def merge_cells_by_column_test(table, column):
    # 初始化合并的第一个cell
    j = 0
    head_cell = table.cell(j, column)
    head_cell_text = head_cell.text
    for i in range(1, len(table.rows)):
        if(head_cell.text == table.cell(i, column).text):
            head_cell.merge(table.cell(i, column))
            head_cell.paragraphs[0].runs[0].font.size = Pt(9)
            head_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            head_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            head_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(69, 69, 69)
            head_cell.text = head_cell_text
        else:
            head_cell = table.cell(i, column)
            head_cell.paragraphs[0].runs[0].font.size = Pt(9)
            head_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            head_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            head_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(69, 69, 69)
            head_cell_text = head_cell.text


def generate_table(file_name, doc):
    df = pd.read_csv('data/{}.csv'.format(file_name))
    table = doc.add_table(1, cols=len(df.columns), style='Table Grid')
    table.autofit = True
    table.columns[0].width = Inches(1)
    table.columns[1].width = Inches(2)
    table.columns[2].width = Inches(1)




    header_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        cell = header_cells[i]
        cell.text = col
        # cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(69, 69, 69)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    table.rows[0].height = Pt(24)
    # 填入数据
    for i, row in enumerate(df.itertuples(), start=1):
        cells = table.add_row().cells
        table.rows[i].height = Pt(24)
        for j, value in enumerate(row[1:], start=0):
            cell = cells[j]
            cell.text = str(value)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(69, 69, 69)


            # cells[j].font.size = Pt(1)

    merge_cells_by_column_test(table, 0)
    merge_cells_by_column_test(table, 1)


    for i in range(len(table.columns)):
        set_table_header_bg_color(table.rows[0].cells[i])
        # table.rows[0].cells[i].border.top.color.rgb = (0,0,0)

        
    # table.style = 'Light Grid'
    table.style.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table.style.font.size = Pt(8)

if __name__ == '__main__':
    document = Document()
    text_style = document.styles['Normal']
    
    text_style.font.name = u'黑体'
    text_style.font.size = Pt(9.5)
    text_style.font.color.rgb = RGBColor(69, 69, 69)
    text_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    head_1_style = document.styles['Heading 1']

    head_1_style.font.name = u'黑体'
    head_1_style.font.size = Pt(15)
    # head_1_style.font.color.rgb = RGBColor(0, 0, 0)
    head_1_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')


    head_2_style = document.styles['Heading 2']

    head_2_style.font.name = u'黑体'
    head_2_style.font.size = Pt(13)
    head_2_style.font.color.rgb = RGBColor(69, 69, 69)
    head_2_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')


    document.sections[0].top_margin = Cm(1.27)
    document.sections[0].bottom_margin = Cm(1.27)
    document.sections[0].left_margin = Cm(1.27)
    document.sections[0].right_margin = Cm(1.27)





    document.add_heading('总结概述', level=1)
    p1 = document.add_paragraph()
    p1.add_run('1. 整体风险').font.size=Pt(10.5)
    draw.draw('1')

    p2 = document.add_paragraph()
    p2.add_run('本次报告分析周期为：(2020年01月01日到2023年02月28日)，通过发票、财务报表、纳税申报表的综合分析，共检测出风险点14项，其中高风险3项，中风险6项，低风险5项')
    document.add_picture('fig/1.png')
    # p1.paragraph_format.line_spacing = Pt(15)
    p3 = document.add_paragraph()
    p3.add_run('2. 具体风险如下').bold = True
    generate_table('1', document)


    document.add_heading('1. 企业基本信息', level=1)
    document.add_heading('1.1 投资方情况', level=2)
    document.add_heading('1.2 纳税信用等级变更', level=2)
    document.add_heading('1.3 税务处罚情况', level=2)


    document.add_heading('2. 发票分析', level=1)
    document.add_heading('2.1 进销对比分析', level=2)
    document.add_heading('2.1.1 近12个月各税率销售、采购额', level=3)
    document.add_heading('2.1.2 近12个月上下游发票税额对比分析', level=3)
    document.add_heading('2.1.3 近12个月互开发票风险', level=3)
    document.add_heading('2.1.4 近12个月购销两头在外风险', level=3)
    document.add_heading('2.1.5 近12个月发票税号与企业名称不匹配', level=3)
    document.add_heading('2.1.6 销售、采购商品明细', level=3)
    document.add_heading('2.2 虚开发票风险', level=2)
    document.add_heading('2.2.1 近12个月公司规模与开票额匹配分析', level=3)
    document.add_heading('2.2.2 红冲、作废、顶额开票分析', level=3)
    document.add_heading('2.2.3 近12个月前十大客户分析', level=3)
    document.add_heading('2.2.4 近12个月前十大行政处罚信息', level=3)
    document.add_heading('2.3 采购虚假发票风险', level=2)
    document.add_heading('2.3.1 近12个月零税额、顶额发票分析', level=3)
    document.add_heading('2.3.2 近12个月前十大供应商分析', level=3)


    document.add_heading('3. 财务涉税风险评估', level=1)
    document.add_heading('3.1 财务涉税风险评估', level=2)
    document.add_heading('3.1.1 销售毛利率', level=3)
    document.add_heading('3.1.2 营业利润率', level=3)
    document.add_heading('3.1.3 财务费用率', level=3)
    document.add_heading('3.1.4 管理费用率', level=3)
    document.add_heading('3.1.5 销售费用率', level=3)
    document.add_heading('3.1.6 研发费用率、委托境外研发比例是否符合高新企业标准分析', level=3)
    document.add_heading('3.1.7 期间费用变动率与营业收入变动率弹性系数分析', level=3)
    document.add_heading('3.1.8 业务招待费占营业收入比值分析', level=3)
    document.add_heading('3.1.9 差旅费占营业收入比值分析', level=3)
    document.add_heading('3.1.10 广告费和业务宣传费明细占比分析', level=3)
    document.add_heading('3.1.11 咨询顾问费明细占比分析', level=3)
    document.add_heading('3.1.12 其他费用明细占比分析', level=3)
    document.add_heading('3.1.13 减值准备对存货占比分析', level=3)
    document.add_heading('3.1.14 固定资产综合折旧率变动异常分析', level=3)
    document.add_heading('3.1.15 超额税前扣除公益性捐赠支出的风险', level=3)
    document.add_heading('3.1.16 未分配利润对实收资本比值过高', level=3)
    document.add_heading('3.2 隐匿收入指标综合分析', level=2)
    document.add_heading('3.2.1 其他应收款变动分析', level=3)
    document.add_heading('3.2.2 其他应付款变动分析', level=3)
    document.add_heading('3.2.3 存货余额、预收账款余额变动分析', level=3)
    document.add_heading('3.3 虚增成本指标综合分析', level=2)
    document.add_heading('3.3.1 应付款项变动分析', level=3)
    document.add_heading('3.3.2 应交税费变动分析', level=3)
    document.add_heading('3.3.3 成本费用、净经营资产异常变动分析', level=3)
    document.add_heading('3.3.4 企业盈利异常', level=3)


    document.add_heading('4. 税务风险评估', level=1)
    document.add_heading('4.1 近三年实缴税额信息', level=2)
    document.add_heading('4.2 增值税', level=2)
    document.add_heading('4.2.1 零申报月数', level=3)
    document.add_heading('4.2.2 增值税税负率（年度）', level=3)
    document.add_heading('4.2.3 增值税税负率（季度）', level=3)
    document.add_heading('4.2.4 增值税农产品收购发票抵扣金额异常（季度）', level=3)
    document.add_heading('4.2.5 有免税收入、简易征税销售额且有进项税额但无进项税额转出（季度）', level=3)
    document.add_heading('4.2.6 增值税留抵退税的风险', level=3)
    document.add_heading('4.3 企业所得税', level=2)
    document.add_heading('4.3.1 企业所得税贡献率', level=3)
    document.add_heading('4.3.2 企业所得税纳税调整增加率', level=3)
    document.add_heading('4.3.3 企业所得税纳税调整减少率', level=3)
    document.add_heading('4.3.4 利润表存在资产减值损失、信用减值损失;申报表不存在纳税调增项', level=3)
    document.add_heading('4.3.5 纳税调整后所得变动率与营业收入变动率弹性系数', level=3)
    document.add_heading('4.3.6 人工费用变动率与营业收入变动率弹性系数', level=3)
    document.add_heading('4.3.7 企业所得税汇缴申报的损失类营业外支出金额', level=3)
    document.add_heading('4.3.8 企业所得税不征税收入调减金额分析', level=3)
    document.add_heading('4.3.9 出资不到位时存在利息支出未纳税调整', level=3)
    document.add_heading('4.3.10 是否符合享受小微企业税收优惠条件检查', level=3)
    document.add_heading('4.3.11 咨询顾问费与发票差异', level=3)
    document.add_heading('4.4 个人所得税', level=2)
    document.add_heading('4.4.1 分配股息红利未扣缴个人所得税风险预警', level=3)
    document.add_heading('4.5 印花税', level=2)
    document.add_heading('4.5.1 印花税变动分析', level=3)
    document.add_heading('4.5.2 印花税税率推算', level=3)
    document.add_heading('4.6 资源税', level=2)
    document.add_heading('4.6.1 资源税变动分析', level=3)
    document.add_heading('4.7 房产税', level=2)
    document.add_heading('4.7.1 房产税变动分析', level=3)
    document.add_heading('4.7.2 存在投资性房地产却无租金收入或未交房产税', level=3)
    document.add_heading('4.8 城建税', level=2)
    document.add_heading('4.8.1 城建税税率异常', level=3)


    document.add_heading('5. 财税票综合风险评估', level=1)
    document.add_heading('5.1 企业所得税、增值税申报收入与财务报表营业收入、销项发票金额对比分析', level=2)
    document.add_heading('5.2 企业所得税利润总额与财务报表利润总额对比分析', level=2)
    document.add_heading('5.3 增值税应纳税额与毛利比值分析', level=2)


    document.add_heading('6. 当前欠税信息', level=1)





    # document.add_page_break()
    docx_file = 'new.docx'
    # html_file = 'new.html'
    document.save(docx_file)
    # d1 = Document(docx=docx_file)
    pdf_file = 'new.pdf'
    docx2pdf.convert(docx_file,pdf_file)
    # pdfkit.from_file(docx_file, pdf_file)


    for paragraph in document.paragraphs:
        if 'Heading' in paragraph.style.name:
            print(paragraph.text)