from docx import Document


def merge_cells_by_column(table, column_index):
    cells = [row.cells[column_index] for row in table.rows]

    start_index = 0
    for i in range(len(cells) - 1):
        if cells[i].text != cells[i + 1].text:
            if start_index != i:
                merge_cells(table, start_index, i, column_index)
            start_index = i + 1

    if start_index != len(cells) - 1:
        merge_cells(table, start_index, len(cells) - 1, column_index)


def merge_cells(table, start_row, end_row, column_index):
    cell_range = table.cell(start_row, column_index), table.cell(end_row, column_index)
    for cell in cell_range:
        cell.width = cell.width + cell_range[1].width
        cell.vertical_alignment = cell_range[0].vertical_alignment
        cell.paragraphs[0].alignment = cell_range[0].paragraphs[0].alignment
    cell_range[1]._element.getparent().remove(cell_range[1]._element)


# 创建一个空白的Word文档
doc = Document()

# 添加一个具有4列和4行的表格
table = doc.add_table(rows=5, cols=4)

# 填入数据
data = [
    ['A', 'B', 'C', 'D'],
    ['A', 'B', 'C', 'D'],
    ['A', 'B', 'C', 'D'],
    ['E', 'F', 'G', 'H'],
    ['E', 'F', 'G', 'H']
]

for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.text = data[i][j]

# 按列合并相同的单元格
merge_cells_by_column(table, 1)  # 合并第二列

# 保存Word文档
doc.save('output.docx')
