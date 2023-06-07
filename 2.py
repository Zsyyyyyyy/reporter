from docx import Document


def merge_cells_by_column(table, column_index):
    cells = [row.cells[column_index] for row in table.rows]

    for i in range(len(cells) - 1):
        if cells[i].text == cells[i + 1].text:
            cells[i + 1].merge(cells[i])


# 创建一个空白的Word文档
doc = Document()

# 添加一个具有4列和4行的表格
table = doc.add_table(rows=5, cols=4)

# 填入数据
data = [
    ['A', 'B', 'C', 'D'],
    ['A', 'B', 'C', 'D'],
    ['A', 'B', 'C', 'D'],
    ['E', 'F', 'G', 'H'],
    ['E', 'F', 'G', 'H']
]

for i, row in enumerate(table.rows):
    for j, cell in enumerate(row.cells):
        cell.text = data[i][j]

# 按列合并单元格
merge_cells_by_column(table, 1)  # 合并第二列

p1 = doc.add_paragraph('Hello, World!')
prior_paragraph = p1.insert_paragraph_before('Lorem ipsum')

# 保存Word文档
doc.save('output.docx')
