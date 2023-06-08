from docx import Document
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
from docx.enum.table import WD_TABLE_ALIGNMENT
import numpy as np
from docx.oxml.shared import qn
from docx.oxml.xmlchemy import OxmlElement


def set_table_header_bg_color(tc):
    """
    set background shading for Header Rows
    """
    tblCellProperties = tc._element.tcPr
    clShading = OxmlElement('w:shd')
    clShading.set(qn('w:fill'), "d3d3d3")  # Hex of Dark Blue Shade {R:0x00, G:0x51, B:0x9E}
    tblCellProperties.append(clShading)


def merge_cells_by_column(table, column_index):
    cells = [row.cells[column_index] for row in table.rows]

    for i in range(len(cells) - 1):
        if cells[i].text == cells[i + 1].text:
            cells[i + 1].merge(cells[i])


def merge_cells_by_column_test(table, column):
    # 初始化合并的第一个cell
    j = 0
    head_cell = table.cell(j, column)
    head_cell_text = head_cell.text
    for i in range(1, len(table.rows)):
        if(head_cell.text == table.cell(i, column).text):
            head_cell.merge(table.cell(i, column))
            head_cell.text = head_cell_text
        else:
            head_cell = table.cell(i, column)
            head_cell_text = head_cell.text


def generate_table(file_name, doc):
    df = pd.read_csv('data/{}.csv'.format(file_name))
    # print('行数' + str(len(df)))
    # print('列数' + str(len(df.columns)))
    # doc = Document(doc)
    table = doc.add_table(1, cols=len(df.columns), style='Table Grid')
    header_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        header_cells[i].text = col

    # 填入数据
    for i, row in enumerate(df.itertuples(), start=1):
        cells = table.add_row().cells
        for j, value in enumerate(row[1:], start=0):
            cells[j].text = str(value)

    merge_cells_by_column_test(table, 0)
    merge_cells_by_column_test(table, 1)

    for i in range(len(table.columns)):
        set_table_header_bg_color(table.rows[0].cells[i])
        table.rows[0].cells[i].width = Inches(1000)
    # table.style = 'Light Grid'
    table.style.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table.style.font.size = Pt(8)


if __name__ == '__main__':
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].font.size = Pt(8)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.add_heading('总结概述', level=1)
    # document.add_heading('1. 整体风险', level=4)
    p1 = document.add_paragraph()
    p1.add_run('1. 整体风险').bold = True
    p2 = document.add_paragraph()
    p2.add_run('本次报告分析周期为：（2020年01月01日到2023年02月28日），通过发票、财务报表、纳税申报表的综合分析，共检测出风险点14项，其中高风险3项，中风险6项，低风险5项')
    # p2 = document.add_paragraph('本次报告分析周期为：（2020年01月01日到2023年02月28日），通过发票、财务报表、纳税申报表的综合分析，共检测出风险点14项，其中高风险3项，中风险6项，低风险5项')
    # p1.paragraph_format.line_spacing = Pt(15)
    # p1.style.font.size = Pt(8)
    # document.add_heading('2. 具体风险如下', level=4)
    # p3 = document.add_paragraph('2. 具体风险如下').bold = True
    p3 = document.add_paragraph()
    p3.add_run('2. 具体风险如下').bold = True
    generate_table('1', document)


    document.add_heading('1. 企业基本信息', level=1)
    document.add_heading('2. 发票分析', level=1)
    document.add_heading('3. 财务涉税风险分析', level=1)





    # document.add_page_break()

    document.save('new.docx')
    